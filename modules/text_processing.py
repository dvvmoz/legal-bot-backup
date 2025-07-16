"""
Модуль для обработки текста и извлечения данных из PDF-файлов и Word документов.
"""
import re
import fitz  # PyMuPDF
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)

def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Извлекает весь текст из PDF-файла.
    
    Args:
        pdf_path: Путь к PDF-файлу
        
    Returns:
        Извлеченный текст из всех страниц PDF
        
    Raises:
        FileNotFoundError: Если файл не найден
        Exception: При ошибке чтения PDF
    """
    try:
        doc = fitz.open(pdf_path)
        full_text = ""
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()  # type: ignore
            full_text += page_text + "\n\n"
            
        doc.close()
        logger.info(f"Извлечено {len(full_text)} символов из файла {pdf_path}")
        return full_text.strip()
        
    except FileNotFoundError:
        logger.error(f"Файл не найден: {pdf_path}")
        raise
    except Exception as e:
        logger.error(f"Ошибка при чтении PDF файла {pdf_path}: {e}")
        raise

def extract_text_from_docx(docx_path: str) -> str:
    """
    Извлекает весь текст из DOCX-файла.
    
    Args:
        docx_path: Путь к DOCX-файлу
        
    Returns:
        Извлеченный текст из документа
        
    Raises:
        FileNotFoundError: Если файл не найден
        Exception: При ошибке чтения DOCX
    """
    try:
        from docx import Document
        
        doc = Document(docx_path)
        full_text = ""
        
        # Извлекаем текст из всех параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text += paragraph.text + "\n"
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text += cell.text + "\n"
        
        logger.info(f"Извлечено {len(full_text)} символов из файла {docx_path}")
        return full_text.strip()
        
    except ImportError:
        logger.error("Для работы с DOCX файлами установите библиотеку: pip install python-docx")
        raise Exception("Библиотека python-docx не установлена")
    except FileNotFoundError:
        logger.error(f"Файл не найден: {docx_path}")
        raise
    except Exception as e:
        logger.error(f"Ошибка при чтении DOCX файла {docx_path}: {e}")
        raise

def extract_text_from_doc(doc_path: str) -> str:
    """
    Извлекает весь текст из DOC-файла (старый формат Word).
    
    Args:
        doc_path: Путь к DOC-файлу
        
    Returns:
        Извлеченный текст из документа
        
    Raises:
        FileNotFoundError: Если файл не найден
        Exception: При ошибке чтения DOC
    """
    try:
        import win32com.client
        
        # Создаем объект Word приложения
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        # Открываем документ
        doc = word.Documents.Open(os.path.abspath(doc_path))
        
        # Извлекаем текст
        full_text = doc.Content.Text
        
        # Закрываем документ и приложение
        doc.Close()
        word.Quit()
        
        logger.info(f"Извлечено {len(full_text)} символов из файла {doc_path}")
        return full_text.strip()
        
    except ImportError:
        logger.error("Для работы с DOC файлами установите библиотеку: pip install pywin32")
        raise Exception("Библиотека pywin32 не установлена или MS Word не найден")
    except FileNotFoundError:
        logger.error(f"Файл не найден: {doc_path}")
        raise
    except Exception as e:
        logger.error(f"Ошибка при чтении DOC файла {doc_path}: {e}")
        raise

def extract_text_from_document(file_path: str) -> str:
    """
    Универсальная функция для извлечения текста из документов различных форматов.
    
    Args:
        file_path: Путь к файлу документа
        
    Returns:
        Извлеченный текст из документа
        
    Raises:
        ValueError: Если формат файла не поддерживается
        Exception: При ошибке чтения файла
    """
    file_extension = Path(file_path).suffix.lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.doc':
        return extract_text_from_doc(file_path)
    else:
        supported_formats = ['.pdf', '.docx', '.doc']
        raise ValueError(f"Неподдерживаемый формат файла: {file_extension}. "
                        f"Поддерживаемые форматы: {', '.join(supported_formats)}")

def get_supported_extensions() -> list[str]:
    """
    Возвращает список поддерживаемых расширений файлов.
    
    Returns:
        Список расширений файлов
    """
    return ['.pdf', '.docx', '.doc']

def is_supported_document(file_path: str) -> bool:
    """
    Проверяет, поддерживается ли формат документа.
    
    Args:
        file_path: Путь к файлу
        
    Returns:
        True если формат поддерживается, False в противном случае
    """
    file_extension = Path(file_path).suffix.lower()
    return file_extension in get_supported_extensions()

def recursive_semantic_splitter(text: str, separators: list[str], max_chunk_size: int) -> list[str]:
    """
    Рекурсивно разделяет текст на семантические части, используя иерархию разделителей.

    Args:
        text (str): Исходный текст для разделения.
        separators (list[str]): Список регулярных выражений для разделителей,
                                отсортированных от самого крупного (например, "Глава")
                                до самого мелкого (например, абзац).
        max_chunk_size (int): Максимальный размер чанка. Если чанк после всех
                              разделений все еще больше, он будет разделен по размеру.

    Returns:
        list[str]: Список текстовых чанков.
    """
    final_chunks = []

    # 1. Если текст уже достаточно мал, возвращаем его как есть.
    if len(text) <= max_chunk_size:
        if text.strip(): # Убедимся, что не добавляем пустые строки
             final_chunks.append(text.strip())
        return final_chunks

    # 2. Если разделители закончились, а текст все еще большой, делим его по размеру.
    if not separators:
        for i in range(0, len(text), max_chunk_size):
            chunk = text[i:i + max_chunk_size]
            if chunk.strip():
                final_chunks.append(chunk.strip())
        return final_chunks

    # 3. Рекурсивный шаг: пытаемся разделить текст текущим разделителем.
    current_separator = separators[0]
    remaining_separators = separators[1:]
    
    # Используем re.split с lookbehind, чтобы сохранить разделитель в начале строки
    try:
        # (?=...) - это positive lookahead, он находит совпадение, но не включает его в результат разделения
        # Это позволяет сохранить заголовок ("Глава 1") в начале следующего чанка.
        chunks = re.split(f'(?={current_separator})', text)
    except re.error as e:
        logger.error(f"Ошибка в регулярном выражении '{current_separator}': {e}")
        # В случае ошибки просто переходим к следующему разделителю
        return recursive_semantic_splitter(text, remaining_separators, max_chunk_size)

    for chunk in chunks:
        if not chunk.strip():
            continue
        
        # Если чанк после разделения все еще слишком большой, рекурсивно вызываем функцию
        # для этого чанка с оставшимися разделителями.
        if len(chunk) > max_chunk_size:
            finer_chunks = recursive_semantic_splitter(chunk, remaining_separators, max_chunk_size)
            final_chunks.extend(finer_chunks)
        else:
            final_chunks.append(chunk.strip())
            
    return final_chunks

def split_text_into_structure(text: str) -> list[str]:
    """
    Разделяет текст на семантические части с использованием рекурсивного алгоритма.
    
    Распознает следующие структурные элементы:
    - Главы: Глава 1., Глава 2.
    - Статьи: Статья 1., Статья 2.
    - Разделы: Раздел 1., Раздел 2.
    - Параграфы: § 1., § 2.
    - Пункты: 1., 2., 10.
    - Подпункты: 1.1., 1.2.3.
    - Буквенные пункты: а), б), в)
    - Абзацы (пустые строки)
    - Нумерованные списки: 1), 2)
    
    Args:
        text: Исходный текст для разделения
        
    Returns:
        Список структурированных текстовых блоков
    """
    if not text or not text.strip():
        return []
    
    # Определяем иерархию разделителей для юридических документов
    # Используем регулярные выражения для большей гибкости
    separators = [
        r"Глава\s*\d+\.",           # Глава 1.
        r"Раздел\s*\d+\.",          # Раздел 1.
        r"Статья\s*\d+\.",          # Статья 1.
        r"§\s*\d+\.",               # § 1.
        r"^\s*\d+\.",               # 1. (нумерованные пункты)
        r"^\s*\d+\.\d+\.",          # 1.1. (подпункты)
        r"^\s*[а-яА-Я]\)\s+",       # а) (буквенные пункты)
        r"\n\s*\n",                 # Разделитель абзацев (пустая строка)
        r"^\s*\d+\)\s+",            # 1) (нумерованные списки)
    ]
    
    # Устанавливаем максимальный размер чанка (можно настроить)
    max_chunk_size = 1000
    
    try:
        chunks = recursive_semantic_splitter(text, separators, max_chunk_size)
        
        # Фильтруем пустые блоки и слишком короткие
        filtered_blocks = []
        for block in chunks:
            clean_block = block.strip()
            if clean_block and len(clean_block) > 10:  # Минимальная длина блока
                filtered_blocks.append(clean_block)
        
        logger.info(f"Разделено на {len(filtered_blocks)} структурированных блоков")
        return filtered_blocks
        
    except Exception as e:
        logger.error(f"Ошибка при семантическом разделении: {e}")
        # В случае ошибки возвращаем простое разделение на абзацы
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip() and len(p.strip()) > 10]
        logger.info(f"Использовано простое разделение на {len(paragraphs)} абзацев")
        return paragraphs

def clean_text(text: str) -> str:
    """
    Очищает текст от лишних пробелов и символов.
    
    Args:
        text: Исходный текст
        
    Returns:
        Очищенный текст
    """
    if not text:
        return ""
    
    # Удаляем лишние пробелы и переводы строк
    cleaned = re.sub(r'\s+', ' ', text)
    
    # Удаляем пробелы в начале и конце
    cleaned = cleaned.strip()
    
    return cleaned


class TextProcessor:
    """
    Класс для обработки текста и извлечения данных из PDF-файлов.
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Инициализация процессора текста.
        
        Args:
            chunk_size: Размер чанка в символах
            chunk_overlap: Перекрытие между чанками в символах
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        logger.info(f"Инициализирован TextProcessor с chunk_size={chunk_size}, chunk_overlap={chunk_overlap}")
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        Извлекает весь текст из PDF-файла.
        
        Args:
            pdf_path: Путь к PDF-файлу
            
        Returns:
            Извлеченный текст из всех страниц PDF
        """
        return extract_text_from_pdf(pdf_path)
    
    def split_text(self, text: str) -> list[str]:
        """
        Разделяет текст на чанки для индексации.
        
        Args:
            text: Исходный текст
            
        Returns:
            Список текстовых чанков
        """
        if not text or not text.strip():
            return []
        
        # Сначала пробуем структурированное разделение
        structured_blocks = split_text_into_structure(text)
        
        # Если структурированное разделение дало результат, используем его
        if structured_blocks and len(structured_blocks) > 1:
            # Объединяем мелкие блоки в чанки нужного размера
            chunks = []
            current_chunk = ""
            
            for block in structured_blocks:
                # Если добавление блока не превысит размер чанка
                if len(current_chunk) + len(block) + 1 <= self.chunk_size:
                    if current_chunk:
                        current_chunk += "\n\n" + block
                    else:
                        current_chunk = block
                else:
                    # Сохраняем текущий чанк
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    
                    # Начинаем новый чанк
                    if len(block) <= self.chunk_size:
                        current_chunk = block
                    else:
                        # Если блок слишком большой, разделяем его
                        chunks.extend(self._split_large_block(block))
                        current_chunk = ""
            
            # Добавляем последний чанк
            if current_chunk:
                chunks.append(current_chunk.strip())
            
            return chunks
        
        # Если структурированное разделение не сработало, используем простое разделение
        return self._simple_split(text)
    
    def _split_large_block(self, block: str) -> list[str]:
        """
        Разделяет большой блок на чанки.
        
        Args:
            block: Блок текста для разделения
            
        Returns:
            Список чанков
        """
        chunks = []
        start = 0
        
        while start < len(block):
            end = start + self.chunk_size
            
            # Если это не последний чанк, ищем подходящее место для разделения
            if end < len(block):
                # Ищем ближайший перевод строки или точку
                for i in range(end, max(start + self.chunk_size - 100, start), -1):
                    if block[i] in '\n.!?':
                        end = i + 1
                        break
            
            chunk = block[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Учитываем перекрытие
            start = end - self.chunk_overlap
            if start < 0:
                start = end
        
        return chunks
    
    def _simple_split(self, text: str) -> list[str]:
        """
        Простое разделение текста на чанки фиксированного размера.
        
        Args:
            text: Исходный текст
            
        Returns:
            Список чанков
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Если это не последний чанк, ищем подходящее место для разделения
            if end < len(text):
                # Ищем ближайший пробел или перевод строки
                for i in range(end, max(start + self.chunk_size - 100, start), -1):
                    if text[i] in ' \n\t':
                        end = i
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Учитываем перекрытие
            start = end - self.chunk_overlap
            if start < 0:
                start = end
        
        logger.info(f"Разделено на {len(chunks)} чанков")
        return chunks
    
    def clean_text(self, text: str) -> str:
        """
        Очищает текст от лишних пробелов и символов.
        
        Args:
            text: Исходный текст
            
        Returns:
            Очищенный текст
        """
        return clean_text(text) 